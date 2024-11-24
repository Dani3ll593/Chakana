import sqlite3
from docx import Document

DATABASE = "database/annotations.db"

def save_annotations():
    conn = sqlite3.connect(DATABASE)
    cursor = conn.cursor()
    cursor.execute("CREATE TABLE IF NOT EXISTS annotations (id INTEGER PRIMARY KEY, text TEXT, comment TEXT)")
    for text, comment in zip(st.session_state['highlighted_text'], st.session_state['comments']):
        cursor.execute("INSERT INTO annotations (text, comment) VALUES (?, ?)", (text, comment))
    conn.commit()
    conn.close()

def export_document():
    doc = Document()
    doc.add_heading("Annotated Document", level=1)
    for text, comment in zip(st.session_state['highlighted_text'], st.session_state['comments']):
        doc.add_paragraph(f"Text: {text}")
        doc.add_paragraph(f"Comment: {comment}")
    doc.save("Annotated_Document.docx")
    st.success("Document exported successfully.")