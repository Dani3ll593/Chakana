import json
import os

class CommentManager:
    def __init__(self):
        self.comments = {}
        self.active_section = None

    def set_active_section(self, section):
        self.active_section = section

    def get_active_section(self):
        return self.active_section

    def add_comment(self, section, comment):
        self.comments[section] = comment

    def delete_comment(self, section):
        if section in self.comments:
            del self.comments[section]

    def get_comments(self):
        return self.comments

    def export_with_comments(self, file_name, content, sections):
        export_path = f"output/{file_name}_with_comments.docx"
        from docx import Document
        doc = Document()
        for i, section in enumerate(sections):
            doc.add_paragraph(section)
            if i in self.comments:
                doc.add_paragraph(f"Comentario: {self.comments[i]}", style="Intense Quote")
        doc.save(export_path)
        return export_path